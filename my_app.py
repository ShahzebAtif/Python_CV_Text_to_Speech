from docx import Document
from docx.shared import Inches
import pyttsx3

#Python Text to Speech
def speak(text):
    pyttsx3.speak(text)

document = Document()

# Profile Picture
document.add_picture('me.png', width=Inches(2.0))

# Name, Phone & Email details
speak('What is your name? ')
name = input('What is your name? ')
speak('Hello ' + name + 'Hope you are doing well today. We will be making CV using Shahzebs Code')

speak('Please enter your Phone Number? ')
phone_number = input('What is your Phone Number? ')
speak('What is your Email? ')
email = input('What is your email? ')

document.add_paragraph(
    name.title() + ' | ' + phone_number + ' | ' + email.lower())

#About Me
document.add_heading('About Me')
speak('Please Tell us about Yourself? ')
document.add_paragraph(input('Tell about Yourself? ').capitalize())

#Education
document.add_heading('Education')
p = document.add_paragraph()

speak('Please enter university name? ')
university = input('Enter University? ')
speak('Please enter the date you started your university? ')
from_date = input('From Date? ')
speak('Please enter the date you graduated from your university? ')
to_date = input('To Date? ')
speak('Please enter your degree name? ')
degree = input('Enter Degree Name? ')
speak('Please enter your Major? ')
major = input('Enter Major of Degree? ')
speak('Please enter your Minor? ')
minor = input('Enter Minor of Degree? ')

p.add_run(university.title() + '\n').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
p.add_run(degree.title() + '\n')
p.add_run(major.capitalize() + ' & ' + minor.capitalize())

#More Universities
while True:
    speak('Do you have more universities to mention? ')
    has_more_universities = input('Do you have more Universities to mention? Yes or No? ')
    if has_more_universities.lower() == 'yes':

        p = document.add_paragraph()

        speak('Please enter university name? ')
        university = input('Enter University? ')
        speak('Please enter the date you started your university? ')
        from_date = input('From Date? ')
        speak('Please enter the date you graduated from your university? ')
        to_date = input('To Date? ')
        speak('Please enter your degree name? ')
        degree = input('Enter Degree Name? ')
        speak('Please enter your Major? ')
        major = input('Enter Major of Degree? ')
        speak('Please enter your Minor? ')
        minor = input('Enter Minor of Degree? ')

        p.add_run(university.title() + '\n').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True
        p.add_run(degree.title() + '\n')
        p.add_run(major.capitalize() + ' & ' + minor.capitalize())

    else:
        break

#Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('Enter your company name? ')
company = input('Enter Company? ')
speak('When did you join' + company + '? ')
from_date = input('From Date? ')
speak('Please enter the date you resigned or Enter current if you are still working here? ')
to_date = input('To Date? ')
p.add_run(company.title() + '\n').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

speak('Describe your experience at ' + company + '? ')
experience_details = input('Describe your experience at ' + company + '? ')
p.add_run(experience_details.capitalize())

#More Experiences
while True:
    speak('Do you have more experiences to mention? Yes or No? ')
    has_more_experiences = input('Do you have more experiences? Yes or No? ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        speak('Enter your company name? ')
        company = input('Enter Company? ')
        speak('When did you join' + company + '? ')
        from_date = input('From Date? ')
        speak('Please enter the date you resigned or Enter current if you are still working here? ')
        to_date = input('To Date? ')
        p.add_run(company.title() + '\n').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        speak('Describe your experience at ' + company + '? ')
        experience_details = input('Describe your experience at ' + company + '? ')
        p.add_run(experience_details.capitalize())

    else:
        break

#Skills
p = document.add_heading('Skills')
p = document.add_paragraph()

speak('Enter Skills? ')
skills = input('Enter Skills? ')
p.add_run(skills.title()).bold = True
p.style = 'List Bullet'

#More Skills
while True:
    speak('Do you have more Skills to mention? Yes or No? ')
    has_more_skills = input('Do you have more Skills to mention? Yes or No? ')
    if has_more_skills.lower() == 'yes':

        p = document.add_paragraph()

        speak('Enter Skills? ')
        skills = input('Enter Skills? ')
        p.add_run(skills.title()).bold = True
        p.style = 'List Bullet'

    else:
        break
    
#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
speak('We have completed making cv. Thank you for using our code and Have a wonderful day. ')
p.text = "CV generated using Shahzeb’s Code in Python =)"

document.save('cv.docx')
